
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.wait import WebDriverWait
from selenium.common.exceptions import *
from selenium.webdriver.support import expected_conditions as condicao_esperada
from selenium.webdriver.common.by import By
from docx import Document
from docx.shared import Cm
from time import sleep
import os



def init_driver():
    # Fonte de opções de switches https://chromium.googlesource.com/chromium/src/+/master/chrome/common/chrome_switches.cc e  https://peter.sh/experiments/chromium-command-line-switches/
    chrome_options = Options()
    '''
    --start-maximized # Inicia maximizado
    --lang=pt-BR # Define o idioma de inicialização, # en-US , pt-BR
    --incognito # Usar o modo anônimo
    --window-size=800,800 # Define a resolução da janela em largura e altura
    --headless # Roda em segundo plano(com a janela fechada)
    --disable-notifications # Desabilita notificações
    --disable-gpu # Desabilita renderização com GPU
    '''
    arguments = ['--lang=pt-BR', '--window-size=1200,1000', '--incognito','--ignore-ssl-errors',
                '--ignore-certificate-errors', '--disable-logging' ]
    for argument in arguments:
        chrome_options.add_argument(argument)
    # Lista de opções experimentais(nem todas estão documentadas) https://chromium.googlesource.com/chromium/src/+/master/chrome/common/pref_names.cc
    # Uso de configurações experimentais
    chrome_options.add_experimental_option('prefs', {
        # Alterar o local padrão de download de arquivos
        'download.default_directory': 'C:\\Users\\igors\\Documentos\\Estudo_python\Dev Aprender\\destrava_dev\\DesafioSemanal-DestravaDev5',
        # notificar o google chrome sobre essa alteração
        'download.directory_upgrade': True,
        # Desabilitar a confirmação de download
        'download.prompt_for_download': False,
        # Desabilitar notificações
        'profile.default_content_setting_values.notifications': 2,
        # Permitir multiplos downloads
        'profile.default_content_setting_values.automatic_downloads': 1,

    })
    # inicializando o webdriver
    driver = webdriver.Chrome(options=chrome_options)

    wait = WebDriverWait(
        driver,
        10,
        poll_frequency=1,
        #https://selenium-python.readthedocs.io/api.html#module-selenium.common.exceptions
        ignored_exceptions=[
            NoSuchElementException,
            ElementNotVisibleException,
            ElementNotSelectableException,
        ])




    return driver , wait

def criar_word():

    doc = Document()
    doc.save('Cotação do Dolar.docx')
    return doc

def escrever_word(cot,data,site, img, doc):

    # Titulo do documento

    doc.add_heading(f'Cotação Atual do Dólar – {cot} ({data})', 0).bold = True

    # Texto 1 do documento
    doc.add_paragraph(f'''O dólar está no valor de {cot}, na data {data}.\n
    Valor cotado no site {site}\n
    Print da cotação atual:
    ''')

    doc.add_picture(img, width=Cm(15))

    # Texto 2 do documento
    doc.add_paragraph('Cotação feita por - Igor Mussalem')

    return doc.save('Cotação do Dolar.docx')

def coletar_dolar(driver):

    cotacao = driver.find_elements(By.XPATH, "//td[@class='Py(10px) Pstart(10px)']")
    cotacao_text = cotacao[3].text
    return cotacao_text

def coletar_data(driver):
    
    data = driver.find_elements(By.XPATH, "//td[@class='Py(10px) Ta(start) Pend(10px)']")
    data_text = data[0].text
    return data_text

def print_site(driver):

    driver.save_screenshot('cotacao.jpg')
    return 'cotacao.jpg'

def abrir_site(site_escolhido, driver):

    driver.get(site_escolhido)

def clicar_dados_historicos(wait):

    wait.until(condicao_esperada.element_to_be_clickable(
    (By.XPATH,"//li[@data-test='HISTORICAL_DATA']"))).click()

def descer_pagina(driver):

    driver.execute_script("window.scrollTo(0, 500);")

def encontrar_iframe(driver):

    iframe = driver.find_element(By.ID, 'guce-inline-consent-iframe')
    return iframe

def mudar_janela(driver, iframe):

    driver.switch_to.frame(iframe)

def clicar_banner(wait):

    wait.until(condicao_esperada.element_to_be_clickable(
    (By.XPATH, "//button[@name='agree']"))).click()

def fechar_janela(driver):

    driver.switch_to.default_content()

def ajustar_cotacao(driver):

    cot = coletar_dolar(driver)
    cot = cot.replace(',', '.')
    cot_num = float(cot)
    cot_num = round(cot_num,2)
    return cot_num
